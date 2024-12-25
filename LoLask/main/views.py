from dal import autocomplete
from .models import Member, Deadline, EventOrg, EventMem, Character
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponseForbidden, HttpResponse
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Border, Side
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from openpyxl.utils import get_column_letter
from docx import Document
from django.contrib.auth import authenticate, login, logout
from .forms import MemberForm, CharacterForm, DeadlineForm, EventOrgForm, EventMemForm, LoginForm
from django.contrib.auth.decorators import login_required, user_passes_test
import datetime


def role_required(role):
    def decorator(view_func):
        def _wrapped_view(request, *args, **kwargs):
            if request.user.is_authenticated and request.user.role == role:
                return view_func(request, *args, **kwargs)
            return HttpResponseForbidden("У вас нет доступа к этой странице.")

        return _wrapped_view

    return decorator


# главная страница
def index(request):
    return render(request, 'main/index/home.html')


def users_login(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user = authenticate(request, username=username, password=password)
            if user is not None:
                login(request, user)
                return redirect('Home')  # Перенаправление на главную страницу после успешного входа
            else:
                form.add_error(None, 'Неверное имя пользователя или пароль.')
    else:
        form = LoginForm()
    return render(request, 'main/index/login.html', {'form': form})


def users_logout(request):
    logout(request)
    return redirect('Home')


# страницы с участниками
@login_required
@role_required('admin')
def members_admin(request):
    members = Member.objects.all()
    return render(request, 'main/index/members_admin.html', {'members': members})


# страницы с персонажами
@login_required
@role_required('admin')
def characters_admin(request):
    male_characters = Character.objects.select_related('Member_ID').filter(Gender='Male')
    female_characters = Character.objects.select_related('Member_ID').filter(Gender='Female')

    return render(request, 'main/index/characters_admin.html', {
        'male_characters': male_characters,
        'female_characters': female_characters
    })


# страница со списком участников
@login_required
@role_required('member')
def characters_member(request):
    male_characters = Character.objects.select_related('Member_ID').filter(Gender='Male')
    female_characters = Character.objects.select_related('Member_ID').filter(Gender='Female')

    return render(request, 'main/index/characters_member.html', {
        'male_characters': male_characters,
        'female_characters': female_characters
    })


# страницы с ивентами
@login_required
@role_required('admin')
def events_admin(request):
    eventsOrg = EventOrg.objects.all()
    eventsMem = EventMem.objects.all()
    return render(request, 'main/index/events_admin.html', {
        'eventsOrg': eventsOrg,
        'eventsMem': eventsMem
    })


# страница ивентов для просматривающих
def events_viewer(request):
    events = EventOrg.objects.all()
    return render(request, 'main/index/events_viewer.html', {
        'events': events
    })


# страницы с дедлайнами для админа и просматривающего
@login_required
@role_required('admin')
def deadline_admin(request):
    male_deadlines = Deadline.objects.filter(Character_ID__Gender='Male')
    female_deadlines = Deadline.objects.filter(Character_ID__Gender='Female')
    return render(request, 'main/index/deadline_admin.html', {
        'male_deadlines': male_deadlines,
        'female_deadlines': female_deadlines
    })


def deadline_viewer(request):
    male_deadlines = Deadline.objects.filter(Character_ID__Gender='Male')
    female_deadlines = Deadline.objects.filter(Character_ID__Gender='Female')
    return render(request, 'main/index/deadline_viewer.html', {
        'male_deadlines': male_deadlines,
        'female_deadlines': female_deadlines
    })


class MemberAutocomplete(autocomplete.Select2QuerySetView):
    def get_queryset(self):
        if not self.request.user.is_authenticated:
            return Member.objects.none()

        qs = Member.objects.all()

        if self.q:
            qs = qs.filter(FCS__icontains=self.q)

        return qs


@login_required
@role_required('admin')
def export_deadline_to_excel(request):
    workbook = Workbook()
    sheet = workbook.active

    # Устанавливаем название листа в соответствии с текущим месяцем и годом
    current_month_year = datetime.datetime.now().strftime("%B %Y")
    sheet.title = current_month_year

    headers = ["Персонаж", "Ответы", "Долги", "Тип дедлайна", "Доп. время"]
    sheet.append(headers)

    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'),
                         bottom=Side(style='thin'))

    # Функция для добавления данных в таблицу
    def add_data_to_sheet(deadlines):
        for deadline in deadlines:
            character = deadline.Character_ID
            row = [
                character.Name,
                deadline.Sub_answer,
                deadline.Miss_answer,
                deadline.Type_of_deadlines,
                deadline.Extra_time
            ]
            sheet.append(row)

            # Окрашиваем ячейки и добавляем границы
            for col in range(1, 6):
                cell = sheet.cell(row=sheet.max_row, column=col)
                cell.border = thin_border

            if deadline.Sub_answer is not None:
                sheet.cell(row=sheet.max_row, column=2).fill = PatternFill(start_color="93c47d", end_color="93c47d",
                                                                           fill_type="solid")

            if deadline.Miss_answer is not None:
                if deadline.Miss_answer < 4:
                    sheet.cell(row=sheet.max_row, column=3).fill = PatternFill(start_color="f6b26b", end_color="f6b26b",
                                                                               fill_type="solid")
                else:
                    sheet.cell(row=sheet.max_row, column=3).fill = PatternFill(start_color="e6916b", end_color="e6916b",
                                                                               fill_type="solid")

    # Разделяем данные по полу
    male_deadlines = Deadline.objects.filter(Character_ID__Gender='Male')
    female_deadlines = Deadline.objects.filter(Character_ID__Gender='Female')

    # Добавляем данные для мужских персонажей
    sheet.append(["Мужские персонажи"])
    add_data_to_sheet(male_deadlines)

    # Добавляем пустую строку между таблицами
    sheet.append([])

    # Добавляем данные для женских персонажей
    sheet.append(["Женские персонажи"])
    add_data_to_sheet(female_deadlines)

    # Автоматическая подстройка ширины столбцов
    for column in sheet.columns:
        max_length = 0
        column = list(column)
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = (max_length + 2)
        sheet.column_dimensions[get_column_letter(column[0].column)].width = adjusted_width

    # Сохранение файла с именем, включающим текущий месяц и год
    current_month = datetime.datetime.now().strftime("%Y-%m")
    filename = f"deadlines_{current_month}.xlsx"
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename={filename}'
    workbook.save(response)
    return response


@login_required
@role_required('admin')
def export_deadline_to_docx(request):
    document = Document()
    document.add_heading('Дедлайны', 0)

    # Функция для добавления данных в таблицу
    def add_data_to_table(deadlines, title):
        document.add_heading(title, level=1)
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Персонаж'
        hdr_cells[1].text = 'Ответы'
        hdr_cells[2].text = 'Долги'
        hdr_cells[3].text = 'Тип дедлайна'
        hdr_cells[4].text = 'Доп. время'

        for deadline in deadlines:
            character = deadline.Character_ID
            row_cells = table.add_row().cells
            row_cells[0].text = character.Name
            row_cells[1].text = str(deadline.Sub_answer) if deadline.Sub_answer is not None else ""
            row_cells[2].text = str(deadline.Miss_answer) if deadline.Miss_answer is not None else ""
            row_cells[3].text = deadline.Type_of_deadlines
            row_cells[4].text = str(deadline.Extra_time) if deadline.Extra_time else ""

            # Окрашиваем ячейки
            if deadline.Sub_answer is not None:
                shading_elm = parse_xml(r'<w:shd {} w:fill="93c47d"/>'.format(nsdecls('w')))
                row_cells[1]._element.get_or_add_tcPr().append(shading_elm)

            if deadline.Miss_answer is not None:
                if deadline.Miss_answer < 4:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="f6b26b"/>'.format(nsdecls('w')))
                else:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="e6916b"/>'.format(nsdecls('w')))
                row_cells[2]._element.get_or_add_tcPr().append(shading_elm)

        # Добавляем границы к таблице
        tbl = table._tbl
        tblBorders = parse_xml(
            r'<w:tblBorders %s><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/><w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>' % nsdecls(
                'w'))
        tbl.tblPr.append(tblBorders)

    # Разделяем данные по полу
    male_deadlines = Deadline.objects.filter(Character_ID__Gender='Male')
    female_deadlines = Deadline.objects.filter(Character_ID__Gender='Female')

    # Добавляем данные для мужских персонажей
    add_data_to_table(male_deadlines, "Мужские персонажи")

    # Добавляем данные для женских персонажей
    add_data_to_table(female_deadlines, "Женские персонажи")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=deadlines.docx'
    document.save(response)
    return response


# добавление участника
@login_required
@role_required('admin')
def add_member(request):
    errors = ''
    if request.method == 'POST':
        form = MemberForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('MembersAdmin')
        else:
            errors = 'Форма была неверной!'
    form = MemberForm()

    data = {
        'form': form,
        'errors': errors,
    }

    return render(request, 'main/index/add_member.html', data)


# редактирование участника
@login_required
@role_required('admin')
def edit_member(request, member_id):
    errors = ''
    member = Member.objects.get(ID=member_id)
    if request.method == 'POST':
        form = MemberForm(request.POST, instance=member)
        if form.is_valid():
            form.save()
            return redirect('MembersAdmin')
        else:
            errors = 'Форма была неверной!'
    else:
        form = MemberForm(instance=member)
    return render(request, 'main/index/edit_member.html', {'form': form, 'errors': errors, 'member': member})


# удаление участника
@login_required
@role_required('admin')
def delete_member(request, member_id):
    member = get_object_or_404(Member, ID=member_id)
    member.delete()
    return redirect('MembersAdmin')


# добавление персонажа
@login_required
@role_required('admin')
def add_character(request):
    errors = ''
    if request.method == 'POST':
        form = CharacterForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('CharactersAdmin')
        else:
            errors = 'Форма была неверной!'
    form = CharacterForm()

    data = {
        'form': form,
        'errors': errors,
    }

    return render(request, 'main/index/add_character.html', data)


# редактирование персонажа
@login_required
@role_required('admin')
def edit_character(request, character_id):
    errors = ''
    character = Character.objects.get(ID=character_id)
    if request.method == 'POST':
        form = CharacterForm(request.POST, instance=character)
        if form.is_valid():
            form.save()
            return redirect('CharactersAdmin')
        else:
            errors = 'Форма была неверной!'
    else:
        form = CharacterForm(instance=character)
    return render(request, 'main/index/edit_character.html', {'form': form, 'errors': errors, 'character': character})


# удаление персонажа
@login_required
@role_required('admin')
def delete_character(request, character_id):
    character = get_object_or_404(Character, ID=character_id)
    character.delete()
    return redirect('CharactersAdmin')


# добавление дедлайна
@login_required
@role_required('admin')
def add_deadline(request):
    errors = ''
    if request.method == 'POST':
        form = DeadlineForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('DeadlineAdmin')
        else:
            errors = 'Форма была неверной!'
    form = DeadlineForm()

    data = {
        'form': form,
        'errors': errors,
    }

    return render(request, 'main/index/add_deadline.html', data)


# редактирование дедлайна
@login_required
@role_required('admin')
def edit_deadline(request, deadline_id):
    errors = ''
    deadline = Deadline.objects.get(ID=deadline_id)
    if request.method == 'POST':
        form = DeadlineForm(request.POST, instance=deadline)
        if form.is_valid():
            form.save()
            return redirect('DeadlineAdmin')
        else:
            errors = 'Форма была неверной!'
    else:
        form = DeadlineForm(instance=deadline)
    return render(request, 'main/index/edit_deadline.html', {'form': form, 'errors': errors, 'deadline': deadline})


# удаление дедлайна
@login_required
@role_required('admin')
def delete_deadline(request, deadline_id):
    deadline = get_object_or_404(Deadline, ID=deadline_id)
    deadline.delete()
    return redirect('DeadlineAdmin')


# добавление ивента
@login_required
@role_required('admin')
def add_event_org(request):
    errors = ''
    if request.method == 'POST':
        form = EventOrgForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('EventsAdmin')
        else:
            errors = 'Форма была неверной!'
    form = EventOrgForm()

    data = {
        'form': form,
        'errors': errors,
    }

    return render(request, 'main/index/add_event_org.html', data)


# редактирование ивента
@login_required
@role_required('admin')
def edit_event_org(request, event_id):
    errors = ''
    event = EventOrg.objects.get(ID=event_id)
    if request.method == 'POST':
        form = EventOrgForm(request.POST, instance=event)
        if form.is_valid():
            form.save()
            return redirect('EventsAdmin')
        else:
            errors = 'Форма была неверной!'
    else:
        form = EventOrgForm(instance=event)
    return render(request, 'main/index/edit_event_org.html', {'form': form, 'errors': errors, 'event': event})


# удаление ивента
@login_required
@role_required('admin')
def delete_event_org(request, event_id):
    event = get_object_or_404(EventOrg, ID=event_id)
    event.delete()
    return redirect('EventsAdmin')


# добавление участника ивента
@login_required
@role_required('admin')
def add_event_mem(request):
    errors = ''
    if request.method == 'POST':
        form = EventMemForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('EventsAdmin')
        else:
            errors = 'Форма была неверной!'
    form = EventMemForm()

    data = {
        'form': form,
        'errors': errors,
    }

    return render(request, 'main/index/add_event_mem.html', data)


# редактирование участника ивента
@login_required
@role_required('admin')
def edit_event_mem(request, event_id, event_mem_id):
    errors = ''
    event_org = EventOrg.objects.get(ID=event_id)

    # Получаем конкретный объект EventMem
    eventmem = get_object_or_404(EventMem, id=event_mem_id, Event_ID=event_org)

    if request.method == 'POST':
        form = EventMemForm(request.POST, instance=eventmem)
        if form.is_valid():
            form.save()
            return redirect('EventsAdmin')
        else:
            errors = 'Форма была неверной!'
    else:
        form = EventMemForm(instance=eventmem)

    return render(request, 'main/index/edit_event_mem.html', {
        'form': form,
        'errors': errors,
        'event_org': event_org,
        'eventmem': eventmem
    })


# удаление участника ивента
@login_required
@role_required('admin')
def delete_event_mem(request, event_id, event_mem_id):
    event_org = get_object_or_404(EventOrg, ID=event_id)
    eventmem = get_object_or_404(EventMem, id=event_mem_id, Event_ID=event_org)
    eventmem.delete()
    return redirect('EventsAdmin')
